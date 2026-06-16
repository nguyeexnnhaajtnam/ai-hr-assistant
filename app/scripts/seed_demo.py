from __future__ import annotations

import argparse
import uuid
from typing import Any

from docx import Document
from sqlalchemy import select

from app.core.config import get_settings
from app.core.database import SessionLocal
from app.models.candidate import Candidate
from app.models.evaluation import Evaluation
from app.models.job import Job
from app.services.scoring_service import ScoringService


DEMO_JOBS: list[dict[str, Any]] = [
    {
        "title": "Demo - Frontend Developer (Angular)",
        "description": (
            "Hiring Angular frontend engineer with TypeScript, RxJS, HTML/CSS, and at least 3 years of experience. "
            "Nice to have Docker and AWS. Domain preference: fintech."
        ),
        "requirements": {
            "title": "Frontend Developer (Angular)",
            "required_skills": [
                {"name": "Angular", "weight": 28},
                {"name": "TypeScript", "weight": 18},
                {"name": "RxJS", "weight": 12},
                {"name": "HTML/CSS", "weight": 12},
            ],
            "nice_to_have_skills": [
                {"name": "Docker", "weight": 5},
                {"name": "AWS", "weight": 5},
            ],
            "min_years_experience": 3,
            "experience_weight": 20,
            "domain": ["fintech"],
            "domain_weight": 10,
        },
        "candidates": [
            {
                "name": "Nguyen Van A",
                "email": "nguyenvana@gmail.com",
                "phone": "0987 654 321",
                "parsed_cv": {
                    "name": "Nguyen Van A",
                    "email": "nguyenvana@gmail.com",
                    "phone": "0987 654 321",
                    "current_title": "Senior Frontend Developer",
                    "summary": (
                        "5.2 years building Angular and TypeScript products for fintech platforms. "
                        "Strong in RxJS, reusable UI architecture, and delivery with cross-functional teams."
                    ),
                    "location": "Hanoi, Vietnam",
                    "experience_years": 5.2,
                    "skills": [
                        {"name": "Angular", "years": 5},
                        {"name": "TypeScript", "years": 5},
                        {"name": "RxJS", "years": 4},
                        {"name": "HTML/CSS", "years": 5},
                        {"name": "Docker", "years": 2},
                        {"name": "JavaScript", "years": 5},
                    ],
                    "domains": ["fintech", "digital banking"],
                    "projects": [
                        "Built onboarding and KYC flows for a fintech lending app.",
                        "Led migration from AngularJS to Angular 16 for digital wallet portal.",
                    ],
                    "experiences": [
                        "Senior Frontend Developer at FinEdge",
                        "Frontend Developer at NeoBank Studio",
                    ],
                    "education": ["Bachelor of Information Technology"],
                },
            },
            {
                "name": "Tran Thi B",
                "email": "tranthib@gmail.com",
                "phone": "0912 345 888",
                "parsed_cv": {
                    "name": "Tran Thi B",
                    "email": "tranthib@gmail.com",
                    "phone": "0912 345 888",
                    "current_title": "Frontend Engineer",
                    "summary": (
                        "3.8 years in Angular web apps with strong UI delivery. Worked in SaaS and e-commerce, "
                        "comfortable with TypeScript and component systems."
                    ),
                    "location": "Da Nang, Vietnam",
                    "experience_years": 3.8,
                    "skills": [
                        {"name": "Angular", "years": 4},
                        {"name": "TypeScript", "years": 4},
                        {"name": "HTML/CSS", "years": 4},
                        {"name": "JavaScript", "years": 4},
                    ],
                    "domains": ["saas", "e-commerce"],
                    "projects": [
                        "Implemented seller dashboard for commerce platform.",
                        "Built design system components shared across product teams.",
                    ],
                    "experiences": [
                        "Frontend Engineer at BrightCart",
                        "UI Developer at AppCanvas",
                    ],
                    "education": ["Bachelor of Computer Science"],
                },
            },
            {
                "name": "Le Van C",
                "email": "levanc@gmail.com",
                "phone": "0904 556 111",
                "parsed_cv": {
                    "name": "Le Van C",
                    "email": "levanc@gmail.com",
                    "phone": "0904 556 111",
                    "current_title": "Web Developer",
                    "summary": (
                        "2.5 years of frontend work with Angular and JavaScript. Interested in moving into product teams "
                        "with stronger engineering practices."
                    ),
                    "location": "Ho Chi Minh City, Vietnam",
                    "experience_years": 2.5,
                    "skills": [
                        {"name": "Angular", "years": 2},
                        {"name": "HTML/CSS", "years": 3},
                        {"name": "JavaScript", "years": 3},
                    ],
                    "domains": ["edtech"],
                    "projects": ["Built student portal and scheduling dashboard."],
                    "experiences": ["Web Developer at SkillSpring"],
                    "education": ["Bachelor of Software Engineering"],
                },
            },
            {
                "name": "Pham Huy D",
                "email": "phamhuyd@gmail.com",
                "phone": "0933 888 222",
                "parsed_cv": {
                    "name": "Pham Huy D",
                    "email": "phamhuyd@gmail.com",
                    "phone": "0933 888 222",
                    "current_title": "Frontend Developer",
                    "summary": (
                        "1.5 years building internal dashboards with React and some Angular exposure. "
                        "Good delivery pace but limited tenure and no fintech context yet."
                    ),
                    "location": "Can Tho, Vietnam",
                    "experience_years": 1.5,
                    "skills": [
                        {"name": "Angular", "years": 1},
                        {"name": "TypeScript", "years": 1},
                        {"name": "HTML/CSS", "years": 2},
                        {"name": "React", "years": 2},
                    ],
                    "domains": ["logistics"],
                    "projects": ["Built fleet monitoring dashboard for logistics operations."],
                    "experiences": ["Frontend Developer at MoveFast"],
                    "education": ["Bachelor of Information Systems"],
                },
            },
            {
                "name": "Hoang Thi E",
                "email": "hoangthie@gmail.com",
                "phone": "0908 112 445",
                "parsed_cv": {
                    "name": "Hoang Thi E",
                    "email": "hoangthie@gmail.com",
                    "phone": "0908 112 445",
                    "current_title": "UI Developer",
                    "summary": (
                        "1.2 years focused on HTML, CSS, and landing page implementation. "
                        "Strong visual polish but limited Angular and TypeScript depth."
                    ),
                    "location": "Hai Phong, Vietnam",
                    "experience_years": 1.2,
                    "skills": [
                        {"name": "HTML/CSS", "years": 3},
                        {"name": "JavaScript", "years": 2},
                    ],
                    "domains": ["retail"],
                    "projects": ["Created promotional microsites and CRM landing pages."],
                    "experiences": ["UI Developer at BrandPort"],
                    "education": ["Multimedia Design Diploma"],
                },
            },
        ],
    },
    {
        "title": "Demo - Product Designer",
        "description": (
            "Need product designer with Figma, UX research, design systems, and 4 years experience. "
            "Nice to have analytics. Domain preference: B2B SaaS."
        ),
        "requirements": {
            "title": "Product Designer",
            "required_skills": [
                {"name": "Figma", "weight": 24},
                {"name": "UX Research", "weight": 16},
                {"name": "Design Systems", "weight": 14},
            ],
            "nice_to_have_skills": [{"name": "Analytics", "weight": 6}],
            "min_years_experience": 4,
            "experience_weight": 25,
            "domain": ["b2b saas"],
            "domain_weight": 15,
        },
        "candidates": [
            {
                "name": "Ngoc Linh",
                "email": "ngoclinh@gmail.com",
                "phone": "0905 112 223",
                "parsed_cv": {
                    "name": "Ngoc Linh",
                    "email": "ngoclinh@gmail.com",
                    "phone": "0905 112 223",
                    "current_title": "Senior Product Designer",
                    "summary": "6 years in SaaS product design with Figma, research ops, and scalable design systems.",
                    "location": "Ho Chi Minh City, Vietnam",
                    "experience_years": 6.0,
                    "skills": [
                        {"name": "Figma", "years": 6},
                        {"name": "UX Research", "years": 5},
                        {"name": "Design Systems", "years": 4},
                        {"name": "Analytics", "years": 2},
                    ],
                    "domains": ["b2b saas", "workflow tools"],
                    "projects": ["Redesigned collaboration suite for operations teams."],
                    "experiences": ["Senior Product Designer at Workgrid"],
                    "education": ["Bachelor of Visual Communication"],
                },
            }
        ],
    },
]


def build_raw_text(parsed_cv: dict[str, Any]) -> str:
    parts: list[str] = []
    for key in ("name", "current_title", "summary", "location"):
        if parsed_cv.get(key):
            parts.append(str(parsed_cv[key]))
    parts.extend(skill["name"] for skill in parsed_cv.get("skills", []) if skill.get("name"))
    parts.extend(parsed_cv.get("projects", []))
    parts.extend(parsed_cv.get("experiences", []))
    parts.extend(parsed_cv.get("education", []))
    return "\n".join(parts)


def build_feedback(job_title: str, candidate_name: str, score: float, recommendation: str, risks: list[str]) -> str:
    label_map = {
        "strong_match": "Ung vien phu hop rat tot voi vi tri nay.",
        "match": "Ung vien phu hop kha tot va nen duoc dua vao shortlist.",
        "weak_match": "Ung vien co mot so diem phu hop nhung can interview ky hon.",
        "not_match": "Ung vien chua phu hop cho vong tiep theo o phase demo nay.",
    }
    opener = label_map.get(recommendation, "Can review them ky hon.")
    risk_text = " ".join(risks[:2]) if risks else "Khong co rui ro lon duoc ghi nhan trong du lieu demo."
    return (
        f"{candidate_name} dang duoc danh gia cho vi tri {job_title}. "
        f"Diem tong hop hien tai la {score:.1f}/100. {opener} "
        f"Luu y de review: {risk_text} Quyet dinh tuyen dung cuoi cung van thuoc ve HR."
    )


def create_demo_file(candidate_id: uuid.UUID, filename: str, parsed_cv: dict[str, Any]) -> str:
    settings = get_settings()
    target_dir = settings.local_upload_dir / "cvs" / str(candidate_id)
    target_dir.mkdir(parents=True, exist_ok=True)
    target_path = target_dir / filename

    document = Document()
    document.add_heading(parsed_cv.get("name") or "Demo Candidate", level=1)
    document.add_paragraph(parsed_cv.get("current_title") or "Candidate")
    document.add_paragraph(parsed_cv.get("summary") or "")
    document.add_heading("Skills", level=2)
    for skill in parsed_cv.get("skills", []):
        document.add_paragraph(
            f"{skill.get('name')} - {skill.get('years')} years" if skill.get("years") else str(skill.get("name")),
            style="List Bullet",
        )
    document.add_heading("Projects", level=2)
    for project in parsed_cv.get("projects", []):
        document.add_paragraph(str(project), style="List Bullet")
    document.save(target_path)
    return str(target_path)


def seed_demo(reset: bool) -> None:
    settings = get_settings()
    settings.local_upload_dir.mkdir(parents=True, exist_ok=True)
    scoring_service = ScoringService()

    with SessionLocal() as session:
        existing_demo_jobs = session.scalars(select(Job).where(Job.title.like("Demo - %"))).all()
        if reset:
            for job in existing_demo_jobs:
                session.delete(job)
            session.commit()
            existing_demo_jobs = []

        if existing_demo_jobs:
            print("Demo seed data already exists. Use --reset to rebuild it.")
            return

        for job_seed in DEMO_JOBS:
            job = Job(
                title=job_seed["title"],
                description=job_seed["description"],
                requirements=job_seed["requirements"],
            )
            session.add(job)
            session.flush()

            for candidate_seed in job_seed["candidates"]:
                parsed_cv = candidate_seed["parsed_cv"]
                evaluation_result = scoring_service.evaluate(job_requirements=job.requirements, parsed_cv=parsed_cv)

                candidate = Candidate(
                    job_id=job.id,
                    name=candidate_seed["name"],
                    email=candidate_seed["email"],
                    phone=candidate_seed["phone"],
                    original_filename=f"{candidate_seed['name'].replace(' ', '_').lower()}.docx",
                    file_path="",
                    raw_text=build_raw_text(parsed_cv),
                    parsed_cv=parsed_cv,
                )
                session.add(candidate)
                session.flush()

                candidate.file_path = create_demo_file(candidate.id, candidate.original_filename, parsed_cv)

                feedback = build_feedback(
                    job.requirements.get("title", job.title),
                    candidate_seed["name"],
                    evaluation_result.overall_score,
                    evaluation_result.recommendation,
                    evaluation_result.risks,
                )

                evaluation = Evaluation(
                    candidate_id=candidate.id,
                    overall_score=evaluation_result.overall_score,
                    recommendation=evaluation_result.recommendation,
                    score_breakdown=evaluation_result.score_breakdown,
                    matched_skills=evaluation_result.matched_skills,
                    missing_skills=evaluation_result.missing_skills,
                    risks=evaluation_result.risks,
                    feedback=feedback,
                    interview_questions=[
                        f"Hay mo ta du an gan nhat lien quan den {job.requirements.get('title', job.title)}.",
                        "Ban da do luong impact cua cong viec minh nhu the nao?",
                        "Neu join team trong 30 ngay dau, ban se uu tien dieu gi?",
                    ],
                )
                session.add(evaluation)

        session.commit()
        print("Demo seed data created successfully.")


def main() -> None:
    parser = argparse.ArgumentParser(description="Seed demo jobs, candidates, and evaluations.")
    parser.add_argument("--reset", action="store_true", help="Delete and recreate demo data.")
    args = parser.parse_args()
    seed_demo(reset=args.reset)


if __name__ == "__main__":
    main()
